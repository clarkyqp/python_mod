import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from datetime import datetime
import os

def create_log_file():
    """创建日志Word文件的核心函数"""
    # 1. 获取UI中用户输入的内容
    custom_content = entry_content.get().strip()  # 获取括号内的自定义内容
    save_dir = entry_path.get().strip()           # 获取用户选择的保存路径
    
    # 2. 验证输入是否完整
    if not custom_content:
        messagebox.showerror("错误", "请输入括号内的自定义内容！")
        return
    if not save_dir or not os.path.exists(save_dir):
        # 尝试创建目录
        try:
            os.makedirs(save_dir, exist_ok=True)
            if not os.path.exists(save_dir):
                raise Exception("无法创建目录")
        except Exception as e:
            messagebox.showerror("错误", f"请选择有效的保存路径！\n{str(e)}")
            return
    
    # 3. 生成当前日期（格式：年-月日，如2025-0903）
    today = datetime.now()
    date_str = today.strftime("%Y-%m%d")  # %Y：4位年份，%m：2位月份，%d：2位日期
    
    # 4. 构建文件名（格式：日志-(自定义内容)-杨清鹏-日期.docx）
    file_name = f"日志-({custom_content})-杨清鹏-{date_str}.docx"
    # 拼接完整的文件保存路径（路径+文件名）
    full_file_path = os.path.join(save_dir, file_name)
    
    # 5. 检查文件是否已存在，避免重复创建
    if os.path.exists(full_file_path):
        confirm = messagebox.askyesno("提示", f"文件「{file_name}」已存在，是否覆盖？")
        if not confirm:
            return
    
    # 6. 创建Word文档并保存
    try:
        # 新建空白Word文档
        doc = Document()
        # 在文档中添加默认内容
        doc.add_heading(f"日志 - {custom_content}", level=1)
        doc.add_paragraph(f"创建日期：{today.strftime('%Y年%m月%d日')}")
        # 保存文档到指定路径
        doc.save(full_file_path)
        # 提示创建成功
        messagebox.showinfo("成功", f"日志文件已创建：\n{full_file_path}")
    except Exception as e:
        # 捕获异常并提示错误信息
        messagebox.showerror("创建失败", f"出错了：{str(e)}")

def select_save_path():
    """打开文件夹选择对话框，让用户选择文件保存路径"""
    # 打开文件夹选择窗口，返回用户选择的文件夹路径
    selected_dir = filedialog.askdirectory(title="选择日志文件保存位置")
    if selected_dir:  # 若用户选择了路径（未取消）
        entry_path.delete(0, tk.END)  # 清空输入框原有内容
        entry_path.insert(0, selected_dir)  # 将选择的路径填入输入框

# ---------------------- 构建UI界面 ----------------------
if __name__ == "__main__":
    # 1. 创建主窗口
    root = tk.Tk()
    root.title("日志文件自动创建工具")  # 窗口标题
    root.geometry("650x200")  # 窗口初始大小（宽x高）
    root.resizable(False, False)  # 禁止窗口拉伸

    # 2. 设置默认路径为当前文件夹
    default_path = os.getcwd()  # 获取当前工作目录作为默认路径
    
    # 3. 创建UI组件（标签、输入框、按钮）并布局
    # （1）自定义内容输入区域
    label_content = ttk.Label(root, text="括号内自定义内容：")
    label_content.grid(row=0, column=0, padx=15, pady=30, sticky="w")  # 网格布局（行0，列0）
    
    entry_content = ttk.Entry(root, width=30)  # 输入框（用于输入自定义内容）
    entry_content.grid(row=0, column=1, padx=10, pady=30, sticky="w")

    # （2）保存路径选择区域
    label_path = ttk.Label(root, text="文件保存位置：")
    label_path.grid(row=1, column=0, padx=15, pady=10, sticky="w")  # 网格布局（行1，列0）
    
    entry_path = ttk.Entry(root, width=50)  # 输入框（显示选择的保存路径）
    entry_path.grid(row=1, column=1, padx=10, pady=10, sticky="w")
    entry_path.insert(0, default_path)  # 设置默认路径
    
    btn_select_path = ttk.Button(root, text="选择路径", command=select_save_path)  # 选择路径按钮
    btn_select_path.grid(row=1, column=2, padx=10, pady=10)

    # （3）创建文件按钮
    btn_create = ttk.Button(root, text="创建日志文件", command=create_log_file, width=20)
    btn_create.grid(row=2, column=1, padx=10, pady=20)

    # 3. 启动UI窗口循环（保持窗口显示）
    root.mainloop()
