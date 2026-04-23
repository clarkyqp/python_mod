import tkinter as tk
from tkinter import filedialog, messagebox
import markdown
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

class MarkdownConverterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Markdown到Word转换器")
        self.root.geometry("600x400")
        self.root.resizable(True, True)
        
        # 设置字体和样式
        self.font = ('Microsoft YaHei', 10)
        
        # 创建UI组件
        self.create_widgets()
        
        # 选中的文件路径
        self.selected_file = None
    
    def create_widgets(self):
        # 标题标签
        title_label = tk.Label(
            self.root, 
            text="Markdown到Word转换器", 
            font=('Microsoft YaHei', 14, 'bold')
        )
        title_label.pack(pady=10)
        
        # 选择文件按钮
        self.select_btn = tk.Button(
            self.root, 
            text="选择Markdown文件", 
            command=self.select_file,
            font=self.font,
            width=20,
            height=2
        )
        self.select_btn.pack(pady=10)
        
        # 文件路径显示
        self.file_path_var = tk.StringVar()
        self.file_path_label = tk.Label(
            self.root, 
            textvariable=self.file_path_var, 
            font=self.font,
            wraplength=500,
            justify=tk.LEFT,
            bg="#f0f0f0",
            padx=10,
            pady=5
        )
        self.file_path_label.pack(fill=tk.X, padx=20, pady=5)
        
        # 转换按钮
        self.convert_btn = tk.Button(
            self.root, 
            text="转换为Word文档", 
            command=self.convert_md_to_docx,
            font=self.font,
            width=20,
            height=2,
            bg="#4CAF50",
            fg="white"
        )
        self.convert_btn.pack(pady=20)
        self.convert_btn.config(state=tk.DISABLED)
        
        # 状态文本框
        self.status_text = tk.Text(
            self.root, 
            height=8, 
            width=70,
            font=self.font,
            wrap=tk.WORD
        )
        self.status_text.pack(padx=20, pady=10, fill=tk.BOTH, expand=True)
        self.status_text.config(state=tk.DISABLED)
    
    def select_file(self):
        """选择Markdown文件"""
        file_path = filedialog.askopenfilename(
            title="选择Markdown文件",
            filetypes=[("Markdown files", "*.md"), ("All files", "*.*")]
        )
        
        if file_path:
            self.selected_file = file_path
            self.file_path_var.set(f"已选择: {file_path}")
            self.convert_btn.config(state=tk.NORMAL)
            self.update_status(f"已选择文件: {file_path}")
    
    def update_status(self, message):
        """更新状态文本框"""
        self.status_text.config(state=tk.NORMAL)
        self.status_text.insert(tk.END, message + "\n")
        self.status_text.see(tk.END)  # 滚动到底部
        self.status_text.config(state=tk.DISABLED)
    
    def convert_md_to_docx(self):
        """将Markdown转换为Word文档"""
        if not self.selected_file:
            messagebox.showerror("错误", "请先选择一个Markdown文件")
            return
        
        try:
            self.update_status("开始转换...")
            
            # 读取Markdown文件
            with open(self.selected_file, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            self.update_status("已读取Markdown内容")
            
            # 创建Word文档
            doc = Document()
            
            # 简单解析Markdown并转换为Word内容
            lines = md_content.split('\n')
            for line in lines:
                line = line.strip()
                
                # 处理标题
                if line.startswith('#'):
                    level = line.count('#')
                    title_text = line.lstrip('#').strip()
                    if level > 3:  # 限制最大标题级别为3
                        level = 3
                    heading = doc.add_heading(title_text, level=level)
                    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    self.update_status(f"添加标题: {title_text} (级别: {level})")
                
                # 处理列表
                elif line.startswith('- ') or line.startswith('* '):
                    para = doc.add_paragraph(line[2:], style='ListBullet')
                    self.update_status(f"添加列表项: {line[2:]}")
                
                # 处理有序列表
                elif re.match(r'^\d+\. ', line):
                    para = doc.add_paragraph(line.split('. ', 1)[1], style='ListNumber')
                    self.update_status(f"添加有序列表项: {line.split('. ', 1)[1]}")
                
                # 处理普通段落
                elif line:
                    para = doc.add_paragraph(line)
                    self.update_status(f"添加段落: {line[:30]}...")
                
                # 处理空行
                else:
                    doc.add_paragraph()
            
            # 保存Word文档
            output_path = self.selected_file.replace('.md', '.docx')
            doc.save(output_path)
            
            self.update_status(f"转换完成！文件已保存至: {output_path}")
            messagebox.showinfo("成功", f"转换完成！\n文件已保存至: {output_path}")
            
        except Exception as e:
            error_msg = f"转换失败: {str(e)}"
            self.update_status(error_msg)
            messagebox.showerror("错误", error_msg)

if __name__ == "__main__":
    # 确保中文显示正常
    root = tk.Tk()
    app = MarkdownConverterApp(root)
    root.mainloop()
